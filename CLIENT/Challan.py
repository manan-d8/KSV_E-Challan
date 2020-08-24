from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from mysql.connector import MySQLConnection, Error
import mysql.connector
import sys
# import comtypes.client

class challan:
	
	def __init__(self,ChallanNo,VehicleNo,date,name,location,VehicleType,Offences,challancount,conn):
		self.document = Document()
		self.pic = "logo.png"
		self.ChallanNo = ChallanNo
		self.VehicleNo = VehicleNo
		self.date = date
		self.Name = name
		self.Location = location
		self.VehicleType = VehicleType
		self.Offence = Offences
		self.conn  = conn
		self.total = 0
		self.challancount = challancount

	def GenerateChallan(self):

		self.headtableE = self.document.add_table(rows=1, cols=1)
		self.row_cellsE = self.headtableE.rows[0].cells
		self.paragraphE = self.row_cellsE[0].paragraphs[0]
		self.paragraphE.add_run("E-challan").bold = True
		self.paragraphE.alignment = WD_ALIGN_PARAGRAPH.CENTER

		self.headtable = self.document.add_table(rows=1, cols=2)
		self.row_cells = self.headtable.rows[0].cells
		self.paragraph = self.row_cells[0].paragraphs[0]
		self.run = self.paragraph.add_run()
		self.run.add_picture(self.pic, width=Inches(3.25))

		# self.paragraph1 = self.row_cells[1].paragraphs[0]
		# self.paragraph1.add_run("E-challan").bold = True
		# self.paragraph1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.paragraph1 = self.row_cells[1].paragraphs[0]
		self.paragraph1.add_run("DATE :"+self.date).bold = True
		self.paragraph1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		self.p = self.document.add_paragraph('')
		self.p.add_run("\t\t\t\t\t\t\t\t\t\t\t\t").underline = True


		self.table = self.document.add_table(rows=5, cols=3)
		self.table.style = 'TableGrid'
		self.table.rows[0].cells[2].merge(self.table.rows[1].cells[2])
		self.table.rows[0].cells[2].merge(self.table.rows[2].cells[2])
		self.table.rows[0].cells[2].merge(self.table.rows[3].cells[2])
		self.table.rows[0].cells[2].merge(self.table.rows[4].cells[2])


		self.paragraph = self.table.rows[0].cells[2].paragraphs[0]
		self.run = self.paragraph.add_run()
		self.run.add_picture('cropped0.jpg', width=Inches(1.25))

		self.cells = self.table.rows[0].cells
		self.cells[0].text = 'Challan No :'
		self.cells[1].text = self.ChallanNo
		self.cells = self.table.rows[1].cells
		self.cells[0].text = 'Vehicle No :'
		self.cells[1].text = self.VehicleNo
		self.cells = self.table.rows[2].cells
		self.cells[0].text = 'Name:'
		self.cells[1].text = self.Name
		self.cells = self.table.rows[3].cells
		self.cells[0].text = 'Location :'
		self.cells[1].text = self.Location
		self.cells = self.table.rows[4].cells
		self.cells[0].text = 'Vehicle Type :'
		self.cells[1].text = self.VehicleType


		self.p = self.document.add_paragraph('')
		self.p.add_run("\t\t\t\t\t\t\t\t\t\t\t\t").underline = True
		#document.add_paragraph('Offence List', style='Intense Quote')

		self.table2 = self.document.add_table(rows=1, cols=2)
		self.table2.style = 'TableGrid'
		self.cells = self.table2.rows[0].cells
		self.cells[0].text = "Offence"
		self.cells[1].text = "Fine"

		self.OffenceIndex = 0
		if not self.challancount == 0:
			self.OffenceIndex = 1


		for i in self.Offence:
			self.cells = self.table2.add_row().cells
			self.cells[0].text = i[0]
			self.cells[1].text = i[self.OffenceIndex]
		
		self.total=0
		self.cells = self.table2.add_row().cells
		
		for x in self.Offence:
			self.total += int(x[self.OffenceIndex+1])

		self.cells[1].text = "Total : "+str(self.total)

		#self.document.add_page_break()

		self.document.save(r'challans/'+str(self.ChallanNo)+'.docx')
		# self.wordToPdf()

	# def wordToPdf(self):
	# 	wdFormatPDF = 17

	# 	in_file = os.path.abspath(r'challans/'+str(self.ChallanNo)+'.docx')
	# 	out_file = os.path.abspath(r'challans/'+str(self.ChallanNo)+'.pdf')

	# 	word = comtypes.client.CreateObject('Word.Application')
	# 	doc = word.Documents.Open(in_file)
	# 	doc.SaveAs(out_file, FileFormat=wdFormatPDF)
	# 	doc.Close()
	# 	word.Quit()

	def read_file(self,filename):
	    with open(filename, 'rb') as f:
	        photo = f.read()
	    return photo

	def AddChallanToDb(self):
	    ex = os.path.isfile("cropped0.jpg")
	    if ex:
	        self.oflist = [x[0]+" ," for x in self.Offence]
	        self.oflist = "".join(self.oflist)
	        data = self.read_file('cropped0.jpg')
	        sqlins = "INSERT INTO challans(ChallanNo,noplate, date, name, location, vehicletype, offences, totalfine , image)\
	        		 VALUES ('"+self.ChallanNo+"', '"+self.VehicleNo+"', '"+self.date+"', '"+self.Name+"', '"+self.Location+"', '"+self.VehicleType+"', '"+str(self.oflist)+"', "+str(self.total)+",%s)"
	        args = (data,)
	        try:
	            cursor=self.conn.cursor()
	            cursor.execute(sqlins, args)
	            self.conn.commit()
	            return True
	        except Error as e:   
	            print("In Error Block | Error is:")                      
	            print(e)
	        print("uploading finished")
	    else:
	        print("picture does not exist")
	        return False

if __name__=="__main__":
	conn = mysql.connector.connect(host="db4free.net",user="ksv1234",passwd="ldrp-itr",database="ksvrto",port=3306)
	ccc = challan("ab2346cndk7890","up12NM4567","12/2/2020","Manan Darji","AkhabarNagar-ahemdabad","M-bkie/bike",(("no helmet","300","500"),("red light jumping","300","500")),0,conn)
	ccc.GenerateChallan()
	ccc.AddChallanToDb()